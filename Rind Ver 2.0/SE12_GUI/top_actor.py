# 劳模
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from PyQt5.QtWebEngineWidgets import QWebEngineView
from docx import Document
from docx.shared import Inches
import os
import sys
sys.path.append("..")
from SE12_Visual import Actor



class TopActor(object):
    def setupUi(self, dialog):
        dialog.setObjectName("dialog")
        dialog.resize(755, 793)
        self.closeWinBtn = QtWidgets.QPushButton(dialog)
        self.closeWinBtn.setGeometry(QtCore.QRect(580, 730, 131, 31))
        self.closeWinBtn.setObjectName("closeWinBtn")
        self.label_picture = QtWidgets.QLabel(dialog)
        self.label_picture.setGeometry(QtCore.QRect(-180, 0, 1150, 850))
        self.label_picture.setStyleSheet("image: url(:/new/back_actor.jpg);")
        self.label_picture.setText("")
        self.label_picture.setObjectName("label_picture")
        self.pushButton = QtWidgets.QPushButton(dialog)
        self.pushButton.setGeometry(QtCore.QRect(410, 120, 93, 28))
        self.pushButton.setObjectName("pushButton")
        self.label_result = QWebEngineView(self)
        self.label_result.setGeometry(QtCore.QRect(50, 200, 640, 480))
        #self.label_result.setText("")
        self.label_result.setObjectName("label_result")
        self.AddTo = QtWidgets.QPushButton(dialog)
        self.AddTo.setGeometry(QtCore.QRect(310, 730, 121, 28))
        self.AddTo.setObjectName("AddTo")
        self.label = QtWidgets.QLabel(dialog)
        self.label.setGeometry(QtCore.QRect(50, 120, 41, 31))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.spinBox = QtWidgets.QSpinBox(dialog)
        self.spinBox.setGeometry(QtCore.QRect(100, 125, 61, 22))
        self.spinBox.setMinimum(2015)
        self.spinBox.setMaximum(2020)
        self.spinBox.setObjectName("spinBox")
        self.spinBox_2 = QtWidgets.QSpinBox(dialog)
        self.spinBox_2.setGeometry(QtCore.QRect(255, 125, 46, 22))
        self.spinBox_2.setMinimum(1)
        self.spinBox_2.setMaximum(20)
        self.spinBox_2.setObjectName("spinBox_2")
        self.label_2 = QtWidgets.QLabel(dialog)
        self.label_2.setGeometry(QtCore.QRect(210, 120, 41, 31))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.label_picture.raise_()
        self.closeWinBtn.raise_()
        self.pushButton.raise_()
        self.label_result.raise_()
        self.AddTo.raise_()
        self.label.raise_()
        self.spinBox.raise_()
        self.spinBox_2.raise_()
        self.label_2.raise_()

        self.retranslateUi(dialog)
        self.closeWinBtn.clicked.connect(dialog.close)
        QtCore.QMetaObject.connectSlotsByName(dialog)

        self.pushButton.clicked.connect(self.on_click)
        self.AddTo.clicked.connect(self.put_into_report)

    def retranslateUi(self, dialog):
        _translate = QtCore.QCoreApplication.translate
        dialog.setWindowTitle(_translate("dialog", "劳模演员"))
        self.closeWinBtn.setText(_translate("dialog", "返回上一级"))
        self.pushButton.setText(_translate("dialog", "生成图表"))
        self.AddTo.setText(_translate("dialog", "加入到数据报表"))
        self.label.setText(_translate("dialog", "年份"))
        self.label_2.setText(_translate("dialog", "排名"))
		
    def on_click(self,dialog):
        Actor(self.spinBox.value(), self.spinBox_2.value())
        #pixMap = QPixmap("Top_actor.jpg").scaled(self.label_result.width(),self.label_result.height())
        #self.label_result.setPixmap(QPixmap(""))
        #self.label_result.setPixmap(pixMap)
        pwd = os.getcwd()
        url_string=os.path.abspath('./SE12_Cache/Actor.html')
        self.label_result.load(QUrl.fromLocalFile(url_string))

    def cut_image(self, dialog):
        size = self.label_result.size()
        img = QtGui.QImage(size, QtGui.QImage.Format_ARGB32)
        painter = QtGui.QPainter(img)
        self.label_result.render(painter)
        self.label_result.render(painter)
        img.save('./SE12_Cache/Top_actor.png')
        painter.end()
    
    def put_into_report(self,dialog):
        try:
            doc = Document('report.docx')
        except Exception:
            #不存在文件则新建
            self.cut_image(dialog)
            doc = Document()
            doc.save('report.docx')
            doc = Document('report.docx')
            #插入图片
            images = './SE12_Cache/Top_actor.png'    # 要插入的图片
            doc.add_paragraph('\n')   # 添加文字
            doc.add_picture(images, width = Inches(6))  # 添加图, 设置宽度
            doc.save('report.docx')     # 保存路径
            print('Add picture to report Succesfully.')
        else:
            self.cut_image(dialog)
            images = './SE12_Cache/Top_actor.png'
            doc.add_paragraph('\n')
            doc.add_picture(images, width = Inches(6))
            doc.save('report.docx')
            print('Add picture to report Succesfully.')

import pictures
