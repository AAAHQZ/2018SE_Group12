# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from PyQt5.QtWebEngineWidgets import QWebEngineView
import os
import sys
sys.path.append("..")
from SE12_Visual import line 

class Line(object):
    def setupUi(self, dialog):
        dialog.setObjectName("dialog")
        dialog.resize(742, 750)
        #dialog.resize(1742, 1750)
        self.closeWinBtn = QtWidgets.QPushButton(dialog)
        self.closeWinBtn.setGeometry(QtCore.QRect(580, 690, 131, 31))
        self.closeWinBtn.setObjectName("closeWinBtn")
        self.label_picture = QtWidgets.QLabel(dialog)
        self.label_picture.setGeometry(QtCore.QRect(-200, 0, 1150, 850))
        self.label_picture.setStyleSheet("image: url(:/new/back_line.jpg);")
        self.label_picture.setText("")
        self.label_picture.setObjectName("label_picture")
        self.spinBox_year = QtWidgets.QSpinBox(dialog)
        self.spinBox_year.setGeometry(QtCore.QRect(110, 120, 61, 22))
        self.spinBox_year.setMinimum(2015)
        self.spinBox_year.setMaximum(2020)
        self.spinBox_year.setObjectName("spinBox_year")
        self.label_texts = QtWidgets.QLabel(dialog)
        self.label_texts.setGeometry(QtCore.QRect(60, 120, 41, 21))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label_texts.setFont(font)
        self.label_texts.setObjectName("label_texts")
        self.pushButton = QtWidgets.QPushButton(dialog)
        self.pushButton.setGeometry(QtCore.QRect(460, 120, 93, 28))
        self.pushButton.setObjectName("pushButton")
        self.AddTo = QtWidgets.QPushButton(dialog)
        self.AddTo.setGeometry(QtCore.QRect(290, 690, 121, 28))
        self.AddTo.setObjectName("AddTo")
        self.spinBox_year_2 = QtWidgets.QSpinBox(dialog)
        self.spinBox_year_2.setGeometry(QtCore.QRect(230, 120, 61, 22))
        self.spinBox_year_2.setMinimum(2015)
        self.spinBox_year_2.setMaximum(2020)
        self.spinBox_year_2.setObjectName("spinBox_year_2")
        self.spinBox_year_3 = QtWidgets.QSpinBox(dialog)
        self.spinBox_year_3.setGeometry(QtCore.QRect(355, 120, 61, 22))
        self.spinBox_year_3.setMinimum(2015)
        self.spinBox_year_3.setMaximum(2020)
        self.spinBox_year_3.setObjectName("spinBox_year_3")
        self.label_texts_2 = QtWidgets.QLabel(dialog)
        self.label_texts_2.setGeometry(QtCore.QRect(190, 120, 41, 21))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label_texts_2.setFont(font)
        self.label_texts_2.setObjectName("label_texts_2")
        self.label_texts_3 = QtWidgets.QLabel(dialog)
        self.label_texts_3.setGeometry(QtCore.QRect(310, 120, 41, 21))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label_texts_3.setFont(font)
        self.label_texts_3.setObjectName("label_texts_3")
        #self.widget = QtWidgets.QWidget(dialog)
        self.label_result = QWebEngineView(self)
        self.label_result.setGeometry(QtCore.QRect(50, 180, 640, 480))

        self.label_result.setObjectName("widget")
        self.label_picture.raise_()
        self.closeWinBtn.raise_()
        self.spinBox_year.raise_()
        self.label_texts.raise_()
        self.pushButton.raise_()
        self.AddTo.raise_()
        self.spinBox_year_2.raise_()
        self.spinBox_year_3.raise_()
        self.label_texts_2.raise_()
        self.label_texts_3.raise_()
        self.label_result.raise_()

        self.retranslateUi(dialog)
        self.closeWinBtn.clicked.connect(dialog.close)
        QtCore.QMetaObject.connectSlotsByName(dialog)

        self.pushButton.clicked.connect(self.on_click)
        self.AddTo.clicked.connect(self.put_into_report)

    def retranslateUi(self, dialog):
        _translate = QtCore.QCoreApplication.translate
        dialog.setWindowTitle(_translate("dialog", "Dialog"))
        self.closeWinBtn.setText(_translate("dialog", "返回上一级"))
        self.label_texts.setText(_translate("dialog", "年份"))
        self.pushButton.setText(_translate("dialog", "生成图表"))
        self.AddTo.setText(_translate("dialog", "加入到数据报表"))
        self.label_texts_2.setText(_translate("dialog", "年份"))
        self.label_texts_3.setText(_translate("dialog", "年份"))
		
    def on_click(self,dialog):
        line(self.spinBox_year.value(), self.spinBox_year_2.value(),self.spinBox_year_3.value())
        pwd = os.getcwd()
        url_string=os.path.abspath('./SE12_Cache/Line.html')

        self.label_result.load(QUrl.fromLocalFile(url_string))

    def cut_image(self, dialog):
        size = self.label_result.size()
        img = QtGui.QImage(size, QtGui.QImage.Format_ARGB32)
        painter = QtGui.QPainter(img)
        self.label_result.render(painter)
        self.label_result.render(painter)
        img.save('./SE12_Cache/Line.png')
        painter.end()

    def put_into_report(self,dialog):
        try:
            doc = Document('report.docx')
        except Exception:
            #不存在文件则新建
            self.cut_image(dialog)
            doc = Document()
            doc.save('report.docx')
            doc = Document('report.docx')
            #插入图片
            images = './SE12_Cache/Line.png'    # 要插入的图片
            doc.add_paragraph('\n')   # 添加文字
            doc.add_picture(images, width = Inches(6))  # 添加图, 设置宽度
            doc.save('report.docx')     # 保存路径
            print('Add picture to report Succesfully.')
        else:
            self.cut_image(dialog)
            images = './SE12_Cache/Line.png'
            doc.add_paragraph('\n')
            doc.add_picture(images, width = Inches(6))
            doc.save('report.docx')
            print('Add picture to report Succesfully.')
    

import pictures
