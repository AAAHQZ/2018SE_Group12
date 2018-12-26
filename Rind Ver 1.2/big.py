# Straight Pie


from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from PyQt5.QtWebEngineWidgets import QWebEngineView
from SE12_Visual import Straight, Pie
from docx import Document
from docx.shared import Inches
import os
import sys

class Big(object):
    def setupUi(self, dialog):
        dialog.setObjectName("dialog")
        dialog.resize(1330, 850)
        self.closeWinBtn = QtWidgets.QPushButton(dialog)
        self.closeWinBtn.setGeometry(QtCore.QRect(1160, 780, 131, 31))
        self.closeWinBtn.setObjectName("closeWinBtn")
        self.spinBox_season = QtWidgets.QSpinBox(dialog)
        self.spinBox_season.setGeometry(QtCore.QRect(640, 60, 46, 22))
        self.spinBox_season.setMinimum(1)
        self.spinBox_season.setMaximum(4)
        self.spinBox_season.setObjectName("spinBox_season")
        self.label_picture = QtWidgets.QLabel(dialog)
        self.label_picture.setGeometry(QtCore.QRect(-50, -110, 1541, 1071))
        self.label_picture.setStyleSheet("image: url(:/new/back_big.jpg);")
        self.label_picture.setText("")
        self.label_picture.setObjectName("label_picture")
        self.spinBox_year = QtWidgets.QSpinBox(dialog)
        self.spinBox_year.setGeometry(QtCore.QRect(530, 60, 61, 22))
        self.spinBox_year.setMinimum(2015)
        self.spinBox_year.setMaximum(2020)
        self.spinBox_year.setObjectName("spinBox_year")
        self.label_texts = QtWidgets.QLabel(dialog)
        self.label_texts.setGeometry(QtCore.QRect(440, 60, 81, 21))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label_texts.setFont(font)
        self.label_texts.setObjectName("label_texts")
        self.pushButton = QtWidgets.QPushButton(dialog)
        self.pushButton.setGeometry(QtCore.QRect(780, 60, 93, 28))
        self.pushButton.setObjectName("pushButton")
        self.label_result1 = QWebEngineView(self)
        self.label_result1.setGeometry(QtCore.QRect(10, 150, 640, 480))
        #self.label_result1.setText("")
        self.label_result1.setObjectName("label_result1")
        self.AddTo = QtWidgets.QPushButton(dialog)
        self.AddTo.setGeometry(QtCore.QRect(260, 700, 121, 28))
        self.AddTo.setObjectName("AddTo")
        self.label_result2 = QWebEngineView(self)
        self.label_result2.setGeometry(QtCore.QRect(680, 150, 640, 480))
        #self.label_result2.setText("")
        self.label_result2.setObjectName("label_result2")
        self.AddTo_2 = QtWidgets.QPushButton(dialog)
        self.AddTo_2.setGeometry(QtCore.QRect(940, 700, 121, 28))
        self.AddTo_2.setObjectName("AddTo_2")
        self.label_texts_2 = QtWidgets.QLabel(dialog)
        self.label_texts_2.setGeometry(QtCore.QRect(594, 60, 41, 21))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label_texts_2.setFont(font)
        self.label_texts_2.setObjectName("label_texts_2")
        self.label_texts_3 = QtWidgets.QLabel(dialog)
        self.label_texts_3.setGeometry(QtCore.QRect(690, 60, 81, 21))
        font = QtGui.QFont()
        font.setFamily("Agency FB")
        font.setPointSize(12)
        self.label_texts_3.setFont(font)
        self.label_texts_3.setObjectName("label_texts_3")
        self.label_picture.raise_()
        self.closeWinBtn.raise_()
        self.spinBox_season.raise_()
        self.spinBox_year.raise_()
        self.label_texts.raise_()
        self.pushButton.raise_()
        self.label_result1.raise_()
        self.AddTo.raise_()
        self.label_result2.raise_()
        self.AddTo_2.raise_()
        self.label_texts_2.raise_()
        self.label_texts_3.raise_()

        self.retranslateUi(dialog)
        self.closeWinBtn.clicked.connect(dialog.close)
        QtCore.QMetaObject.connectSlotsByName(dialog)

        self.pushButton.clicked.connect(self.on_click)
        self.AddTo.clicked.connect(self.put_into_report1)
        self.AddTo_2.clicked.connect(self.put_into_report2)

    def retranslateUi(self, dialog):
        _translate = QtCore.QCoreApplication.translate
        dialog.setWindowTitle(_translate("dialog", "票房份额"))
        self.closeWinBtn.setText(_translate("dialog", "返回上一级"))
        self.label_texts.setText(_translate("dialog", "选择时间"))
        self.pushButton.setText(_translate("dialog", "生成图表"))
        self.AddTo.setText(_translate("dialog", "加入到数据报表"))
        self.AddTo_2.setText(_translate("dialog", "加入到数据报表"))
        self.label_texts_2.setText(_translate("dialog", "年第"))
        self.label_texts_3.setText(_translate("dialog", "季度"))

    def on_click(self,dialog):
        Straight(self.spinBox_year.value(), self.spinBox_season.value())
        Pie(self.spinBox_year.value(), self.spinBox_season.value())

        pwd = os.getcwd()
        url_string1=os.path.abspath('./SE12_Cache/Straight.html')
        self.label_result1.load(QUrl.fromLocalFile(url_string1))

        url_string2=os.path.abspath('./SE12_Cache/Pie.html')
        self.label_result2.load(QUrl.fromLocalFile(url_string2))

    def cut_image1(self, dialog):
        size = self.label_result1.size()
        img = QtGui.QImage(size, QtGui.QImage.Format_ARGB32)
        painter = QtGui.QPainter(img)
        self.label_result1.render(painter)
        self.label_result1.render(painter)
        img.save('Straight.png')
        painter.end()

    def cut_image2(self, dialog):
        size = self.label_result2.size()
        img = QtGui.QImage(size, QtGui.QImage.Format_ARGB32)
        painter = QtGui.QPainter(img)
        self.label_result2.render(painter)
        self.label_result2.render(painter)
        img.save('Pie.png')
        painter.end()
    
    def put_into_report1(self,dialog):
        try:
            doc = Document('report.docx')
        except Exception:
            #不存在文件则新建
            self.cut_image1(dialog)
            doc = Document()
            doc.save('report.docx')
            doc = Document('report.docx')
            #插入图片
            images = 'Straight.png'    # 要插入的图片
            doc.add_paragraph('\n')   # 添加文字
            doc.add_picture(images, width = Inches(6))  # 添加图, 设置宽度
            doc.save('report.docx')     # 保存路径
            print('Add picture to report Succesfully.')
        else:
            self.cut_image1(dialog)
            images = 'Straight.png'
            doc.add_paragraph('\n')
            doc.add_picture(images, width = Inches(6))
            doc.save('report.docx')
            print('Add picture to report Succesfully.')

    def put_into_report2(self,dialog):
        try:
            doc = Document('report.docx')
        except Exception:
            #不存在文件则新建
            self.cut_image2(dialog)
            doc = Document()
            doc.save('report.docx')
            doc = Document('report.docx')
            #插入图片
            images = 'Pie.png'    # 要插入的图片
            doc.add_paragraph('\n')   # 添加文字
            doc.add_picture(images, width = Inches(6))  # 添加图, 设置宽度
            doc.save('report.docx')     # 保存路径
            print('Add picture to report Succesfully.')
        else:
            self.cut_image2(dialog)
            images = 'Pie.png'
            doc.add_paragraph('\n')
            doc.add_picture(images, width = Inches(6))
            doc.save('report.docx')
            print('Add picture to report Succesfully.')
    
import pictures
